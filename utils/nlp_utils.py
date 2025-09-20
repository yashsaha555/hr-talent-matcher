"""
NLP Utilities for Resume Analysis and Skill Extraction
"""

import re
import nltk
import spacy
from textblob import TextBlob
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import json
from docx import Document
import PyPDF2
from io import BytesIO

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords')

try:
    nltk.data.find('taggers/averaged_perceptron_tagger')
except LookupError:
    nltk.download('averaged_perceptron_tagger')

class ResumeAnalyzer:
    """Advanced NLP-based resume analysis and skill extraction"""
    
    def __init__(self):
        # Load spaCy model (install with: python -m spacy download en_core_web_sm)
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except OSError:
            print("Warning: spaCy model not found. Some features may be limited.")
            self.nlp = None
        
        # Predefined skill categories and keywords
        self.skill_patterns = {
            'Programming Languages': [
                'python', 'java', 'javascript', 'c++', 'c#', 'ruby', 'php', 'go', 'rust',
                'swift', 'kotlin', 'scala', 'r', 'matlab', 'sql', 'html', 'css', 'typescript'
            ],
            'Web Technologies': [
                'react', 'angular', 'vue.js', 'node.js', 'express', 'django', 'flask',
                'spring', 'laravel', 'rails', 'asp.net', 'bootstrap', 'jquery'
            ],
            'Databases': [
                'mysql', 'postgresql', 'mongodb', 'redis', 'oracle', 'sqlite',
                'cassandra', 'elasticsearch', 'neo4j', 'dynamodb'
            ],
            'Cloud Platforms': [
                'aws', 'azure', 'gcp', 'google cloud', 'docker', 'kubernetes',
                'terraform', 'jenkins', 'gitlab', 'github actions'
            ],
            'Data Science': [
                'machine learning', 'deep learning', 'data analysis', 'statistics',
                'pandas', 'numpy', 'scikit-learn', 'tensorflow', 'pytorch', 'keras'
            ],
            'Soft Skills': [
                'leadership', 'communication', 'teamwork', 'problem solving',
                'project management', 'agile', 'scrum', 'mentoring', 'collaboration'
            ],
            'Tools': [
                'git', 'jira', 'confluence', 'slack', 'trello', 'asana',
                'visual studio', 'intellij', 'eclipse', 'postman'
            ]
        }
        
        self.tfidf_vectorizer = TfidfVectorizer(max_features=1000, stop_words='english')
    
    def extract_text_from_pdf(self, pdf_file):
        """Extract text from PDF file"""
        try:
            if isinstance(pdf_file, str):
                with open(pdf_file, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page in pdf_reader.pages:
                        text += page.extract_text()
            else:
                pdf_reader = PyPDF2.PdfReader(BytesIO(pdf_file.read()))
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text()
            return text
        except Exception as e:
            print(f"Error extracting PDF text: {e}")
            return ""
    
    def extract_text_from_docx(self, docx_file):
        """Extract text from DOCX file"""
        try:
            if isinstance(docx_file, str):
                doc = Document(docx_file)
            else:
                doc = Document(BytesIO(docx_file.read()))
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            print(f"Error extracting DOCX text: {e}")
            return ""
    
    def preprocess_text(self, text):
        """Clean and preprocess text"""
        # Convert to lowercase
        text = text.lower()
        
        # Remove special characters but keep spaces
        text = re.sub(r'[^a-zA-Z0-9\s\.\+\#]', ' ', text)
        
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text).strip()
        
        return text
    
    def extract_skills_from_text(self, text):
        """Extract skills from text using pattern matching and NLP"""
        text = self.preprocess_text(text)
        extracted_skills = {}
        
        # Pattern-based skill extraction
        for category, skills in self.skill_patterns.items():
            found_skills = []
            for skill in skills:
                # Use word boundaries to avoid partial matches
                pattern = r'\b' + re.escape(skill.lower()) + r'\b'
                if re.search(pattern, text):
                    found_skills.append(skill)
            
            if found_skills:
                extracted_skills[category] = found_skills
        
        # NLP-based skill extraction using spaCy
        if self.nlp:
            doc = self.nlp(text)
            
            # Extract technical terms and proper nouns that might be skills
            technical_terms = []
            for token in doc:
                if (token.pos_ in ['NOUN', 'PROPN'] and 
                    len(token.text) > 2 and 
                    not token.is_stop and 
                    not token.is_punct):
                    technical_terms.append(token.text.lower())
            
            # Filter technical terms that might be skills
            potential_skills = [term for term in technical_terms 
                              if any(keyword in term for keyword in 
                                   ['tech', 'soft', 'program', 'develop', 'manage', 'analy'])]
            
            if potential_skills:
                extracted_skills['Identified Skills'] = list(set(potential_skills))
        
        return extracted_skills
    
    def extract_experience_years(self, text):
        """Extract years of experience from text"""
        text = text.lower()
        
        # Pattern for "X years of experience"
        patterns = [
            r'(\d+)\s*(?:\+)?\s*years?\s*(?:of\s*)?experience',
            r'(\d+)\s*(?:\+)?\s*years?\s*(?:in|with)',
            r'experience\s*(?:of\s*)?(\d+)\s*(?:\+)?\s*years?',
            r'(\d+)\s*(?:\+)?\s*yrs?\s*(?:of\s*)?(?:exp|experience)'
        ]
        
        years = []
        for pattern in patterns:
            matches = re.findall(pattern, text)
            years.extend([int(match) for match in matches])
        
        # Return the maximum years found, or 0 if none
        return max(years) if years else 0
    
    def extract_education(self, text):
        """Extract education information from text"""
        text = text.lower()
        
        education_patterns = {
            'degree_types': [
                'bachelor', 'master', 'phd', 'doctorate', 'associate',
                'b.s.', 'b.a.', 'm.s.', 'm.a.', 'mba', 'b.tech', 'm.tech'
            ],
            'fields': [
                'computer science', 'engineering', 'business', 'mathematics',
                'physics', 'chemistry', 'biology', 'economics', 'finance'
            ]
        }
        
        found_education = []
        
        for degree in education_patterns['degree_types']:
            if degree in text:
                # Try to find the field associated with this degree
                for field in education_patterns['fields']:
                    if field in text:
                        found_education.append(f"{degree} in {field}")
                        break
                else:
                    found_education.append(degree)
        
        return list(set(found_education))
    
    def extract_certifications(self, text):
        """Extract certifications from text"""
        text = text.lower()
        
        cert_patterns = [
            'certified', 'certification', 'certificate', 'aws', 'azure', 'gcp',
            'pmp', 'scrum master', 'agile', 'cissp', 'cisa', 'cism'
        ]
        
        certifications = []
        for pattern in cert_patterns:
            if pattern in text:
                # Extract surrounding context
                matches = re.finditer(r'\b' + re.escape(pattern) + r'\b', text)
                for match in matches:
                    start = max(0, match.start() - 20)
                    end = min(len(text), match.end() + 20)
                    context = text[start:end].strip()
                    certifications.append(context)
        
        return list(set(certifications))
    
    def analyze_resume(self, text):
        """Comprehensive resume analysis"""
        analysis = {
            'skills': self.extract_skills_from_text(text),
            'experience_years': self.extract_experience_years(text),
            'education': self.extract_education(text),
            'certifications': self.extract_certifications(text),
            'text_length': len(text),
            'word_count': len(text.split()),
            'sentiment': self.analyze_sentiment(text)
        }
        
        return analysis
    
    def analyze_sentiment(self, text):
        """Analyze sentiment of the text"""
        try:
            blob = TextBlob(text)
            return {
                'polarity': blob.sentiment.polarity,
                'subjectivity': blob.sentiment.subjectivity
            }
        except:
            return {'polarity': 0.0, 'subjectivity': 0.0}
    
    def calculate_text_similarity(self, text1, text2):
        """Calculate semantic similarity between two texts"""
        try:
            texts = [self.preprocess_text(text1), self.preprocess_text(text2)]
            tfidf_matrix = self.tfidf_vectorizer.fit_transform(texts)
            similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
            return similarity
        except:
            return 0.0
    
    def extract_contact_info(self, text):
        """Extract contact information from text"""
        contact_info = {}
        
        # Email pattern
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        if emails:
            contact_info['email'] = emails[0]
        
        # Phone pattern
        phone_pattern = r'(\+?1?[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})'
        phones = re.findall(phone_pattern, text)
        if phones:
            contact_info['phone'] = ''.join(phones[0])
        
        # LinkedIn pattern
        linkedin_pattern = r'linkedin\.com/in/[\w-]+'
        linkedin = re.findall(linkedin_pattern, text.lower())
        if linkedin:
            contact_info['linkedin'] = linkedin[0]
        
        # GitHub pattern
        github_pattern = r'github\.com/[\w-]+'
        github = re.findall(github_pattern, text.lower())
        if github:
            contact_info['github'] = github[0]
        
        return contact_info

class SkillMatcher:
    """Advanced skill matching and gap analysis"""
    
    def __init__(self):
        self.tfidf_vectorizer = TfidfVectorizer(max_features=500, stop_words='english')
    
    def calculate_skill_similarity(self, skill1, skill2):
        """Calculate similarity between two skills"""
        try:
            skills = [skill1.lower(), skill2.lower()]
            tfidf_matrix = self.tfidf_vectorizer.fit_transform(skills)
            similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
            return similarity
        except:
            return 0.0
    
    def find_skill_gaps(self, employee_skills, required_skills):
        """Identify skill gaps and suggest improvements"""
        gaps = []
        matches = []
        
        for req_skill, req_level in required_skills.items():
            best_match = None
            best_similarity = 0.0
            
            for emp_skill, emp_level in employee_skills.items():
                similarity = self.calculate_skill_similarity(req_skill, emp_skill)
                if similarity > best_similarity:
                    best_similarity = similarity
                    best_match = (emp_skill, emp_level)
            
            if best_similarity > 0.7:  # Good match threshold
                if best_match[1] >= req_level:
                    matches.append({
                        'required_skill': req_skill,
                        'employee_skill': best_match[0],
                        'required_level': req_level,
                        'employee_level': best_match[1],
                        'similarity': best_similarity,
                        'status': 'meets_requirement'
                    })
                else:
                    gaps.append({
                        'required_skill': req_skill,
                        'employee_skill': best_match[0],
                        'required_level': req_level,
                        'employee_level': best_match[1],
                        'similarity': best_similarity,
                        'gap': req_level - best_match[1],
                        'status': 'needs_improvement'
                    })
            else:
                gaps.append({
                    'required_skill': req_skill,
                    'employee_skill': None,
                    'required_level': req_level,
                    'employee_level': 0,
                    'similarity': 0.0,
                    'gap': req_level,
                    'status': 'missing_skill'
                })
        
        return {
            'gaps': gaps,
            'matches': matches,
            'gap_score': len(gaps) / len(required_skills) if required_skills else 0,
            'match_score': len(matches) / len(required_skills) if required_skills else 0
        }
